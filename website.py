from flask import Flask, request, send_file, render_template_string
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from twilio.rest import Client
from dotenv import load_dotenv
from flask import send_from_directory
import os

app = Flask(__name__)

load_dotenv()
TWILIO_ACCOUNT_SID = os.getenv("TWILIO_ACCOUNT_SID")
TWILIO_AUTH_TOKEN = os.getenv("TWILIO_AUTH_TOKEN")
TWILIO_WHATSAPP_NUMBER = os.getenv("TWILIO_WHATSAPP_NUMBER")  # e.g., whatsapp:+14155238886
TARGET_WHATSAPP_NUMBER = os.getenv("TARGET_WHATSAPP_NUMBER")  # your phone number with whatsapp:

client = Client(TWILIO_ACCOUNT_SID, TWILIO_AUTH_TOKEN)

form_html = '''<!DOCTYPE html>
<html>
<head>
    <title>Admission Form</title>
    <style>
        body { font-family: Arial, sans-serif; background: #f0f4ff; padding: 20px; }
        h2 { color: #1f4e79; text-align: center; }
        form {
            background: white; border-radius: 10px; padding: 20px;
            max-width: 700px; margin: auto;
            box-shadow: 0 0 15px rgba(0,0,0,0.1);
        }
        input, select, textarea {
            width: 100%; padding: 8px; margin: 10px 0;
            border: 1px solid #ccc; border-radius: 5px;
        }
        button {
            background: #1f4e79; color: white; padding: 10px 15px;
            border: none; border-radius: 5px; cursor: pointer;
        }
        button:hover { background: #163a5f; }
        label { font-weight: bold; display: block; margin-top: 10px; }
    </style>
    <script>
        function validateForm() {
            const form = document.forms[0];
            const email = form.Email.value;
            const phonePattern = /^[0-9]{10}$/;
            const emailPattern = /^[^\\s@]+@[^\\s@]+\\.[^\\s@]+$/;
            const dob = new Date(form.DOB.value);
            const age = parseInt(form.Age.value);
            const today = new Date();
            const calculatedAge = today.getFullYear() - dob.getFullYear();
            const ageDiff = Math.abs(calculatedAge - age);

            if (!emailPattern.test(email)) {
                alert("Invalid email format");
                return false;
            }

            if (!phonePattern.test(form.FatherPhone.value) || !phonePattern.test(form.MotherPhone.value)) {
                alert("Phone number must be 10 digits");
                return false;
            }

            if (ageDiff > 1) {
                alert("Age does not match with Date of Birth");
                return false;
            }

            for (const element of form.elements) {
                if (element.hasAttribute("required") && !element.value.trim()) {
                    alert("All fields are required");
                    return false;
                }
            }
            return true;
        }
    </script>
</head>
<body>
    <h2>📄 Darussalam Public School Admission Form</h2>
    <form method="POST" onsubmit="return validateForm();">
        <input name="FullName" placeholder="Full Name" required>
        <input name="DOB" type="date" required>
        <input name="Age" placeholder="Age" required>
        <select name="Gender" required>
            <option value="">Select Gender</option>
            <option>Male</option>
            <option>Female</option>
        </select>
        <input name="Class" placeholder="Class Seeking Admission" required>
        <input name="MotherTongue" placeholder="Mother Tongue" required>
        <select name="Nationality" required>
            <option value="">Select Nationality</option>
            <option value="India">India</option>
            <option value="United States">United States</option>
            <option value="United Kingdom">United Kingdom</option>
            <option value="Canada">Canada</option>
            <option value="Australia">Australia</option>
            <option value="Germany">Germany</option>
            <option value="France">France</option>
            <option value="Italy">Italy</option>
            <option value="Spain">Spain</option>
            <option value="Brazil">Brazil</option>
            <option value="Mexico">Mexico</option>
            <option value="Russia">Russia</option>
            <option value="China">China</option>
            <option value="Japan">Japan</option>
            <option value="South Korea">South Korea</option>
            <option value="Singapore">Singapore</option>
            <option value="Malaysia">Malaysia</option>
            <option value="Thailand">Thailand</option>
            <option value="UAE">United Arab Emirates</option>
            <option value="Saudi Arabia">Saudi Arabia</option>
            <option value="South Africa">South Africa</option>
            <option value="Nigeria">Nigeria</option>
            <option value="Egypt">Egypt</option>
            <option value="Turkey">Turkey</option>
            <option value="Netherlands">Netherlands</option>
            <option value="Switzerland">Switzerland</option>
            <option value="Sweden">Sweden</option>
            <option value="Norway">Norway</option>
            <option value="Denmark">Denmark</option>
            <option value="New Zealand">New Zealand</option>
            <option value="Argentina">Argentina</option>
            <option value="Chile">Chile</option>
            <option value="Colombia">Colombia</option>
            <option value="Pakistan">Pakistan</option>
            <option value="Bangladesh">Bangladesh</option>
            <option value="Nepal">Nepal</option>
            <option value="Sri Lanka">Sri Lanka</option>
            <option value="Qatar">Qatar</option>
            <option value="Other">Other</option>
        </select>
        <select name="Religion" required>
            <option value="">Select Religion</option>
            <option value="Hinduism">Hinduism</option>
            <option value="Islam">Islam</option>
            <option value="Christianity">Christianity</option>
            <option value="Sikhism">Sikhism</option>
            <option value="Buddhism">Buddhism</option>
            <option value="Jainism">Jainism</option>
            <option value="Judaism">Judaism</option>
            <option value="Zoroastrianism">Zoroastrianism</option>
            <option value="Baháʼí Faith">Baháʼí Faith</option>
            <option value="Other">Other</option>
        </select>
        <input name="Caste" placeholder="Caste" required>
        <select name="Category" required>
            <option value="">Select Category</option>
            <option value="SC">SC</option>
            <option value="ST">ST</option>
            <option value="OBC">OBC</option>
            <option value="General">General</option>
        </select>
        <input name="Aadhar" placeholder="Aadhar Number" required>
        <input name="PreviousSchool" placeholder="Previous School" required>
        <input name="FatherName" placeholder="Father's Name" required>
        <input name="FatherPhone" placeholder="Father's Phone Number" required>
        <input name="FatherEduOcc" placeholder="Father's Qualification & Occupation" required>
        <input name="MotherName" placeholder="Mother's Name" required>
        <input name="MotherPhone" placeholder="Mother's Phone Number" required>
        <input name="MotherEduOcc" placeholder="Mother's Qualification & Occupation" required>
        <textarea name="Address" placeholder="Full Address with PIN" required rows="3"></textarea>
        <input name="Email" type="email" placeholder="Email ID" required>
        <input name="PhysicalChallenges" placeholder="Physical Challenges (if any)" required>
        <input name="SiblingDetails" placeholder="Sibling Details (if any)" required>
        <select name="Transport" required>
            <option value="">Need transportation?</option>
            <option>Yes</option>
            <option>No</option>
        </select>
        <div style="text-align: right; clear: both; margin-top: 10px;">
            <button type="submit" style="margin: 0;">Submit</button>
        </div>
    </form>
</body>
</html>
'''

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        form_data = {key: request.form[key] for key in request.form}
        docx_filename = generate_filled_docx(form_data)
        
        # Step 1: Upload your docx file to a static URL or pre-upload to a cloud (you need to implement it)
        # For now assume it is uploaded and publicly accessible via this placeholder:
        media_url = f"https://schooldpsv1.onrender.com/uploads/{docx_filename}"
        
        send_whatsapp_message(media_url)

         # Delete the sent file immediately after sending
        try:
            os.remove(os.path.join("uploads", docx_filename))
        except FileNotFoundError:
            pass

        return "Form submitted and file sent via WhatsApp!"

    return render_template_string(form_html)

@app.route('/uploads/<path:filename>')    
def serve_upload(filename):
    return send_from_directory('uploads', filename)


def set_row_height(row, height_pt):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_pt * 20)))  # height in twips
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)

def generate_admission_filename(user_name, doc):
    os.makedirs("uploads", exist_ok=True)

    safe_name = user_name.strip().replace(" ", "_")

    existing_files = sorted([
        f for f in os.listdir("uploads")
        if f.startswith(safe_name + "_admission") and f.endswith(".docx")
    ], key=lambda x: os.path.getctime(os.path.join("uploads", x)))

    if len(existing_files) >= 5:
        for old_file in existing_files[:len(existing_files) - 4]:
            os.remove(os.path.join("uploads", old_file))

    max_num = 0
    for f in existing_files:
        try:
            number_part = f[len(safe_name) + len("_admission"):-5]
            number = int(number_part)
            max_num = max(max_num, number)
        except ValueError:
            continue

    new_number = max_num + 1
    filename = f"{safe_name}_admission{new_number}.docx"
    output_path = os.path.join("uploads", filename)
    doc.save(output_path)

    return filename

def generate_filled_docx(data):
    template_path = "static/template1.docx"
    doc = Document(template_path)

    doc.add_paragraph("\n")  # Add space

    # Table with headers
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Value'

    # Add data rows
    for key, value in data.items():
        row = table.add_row()
        set_row_height(row, 20)  # set row height to 20pt
        cells = row.cells
        cells[0].text = key
        cells[1].text = value

    # Pad to bottom
    min_rows = 30
    # for _ in range(min_rows - len(data)):
    #     row = table.add_row()
    #     set_row_height(row, 20)
    #     cells = row.cells
    #     cells[0].text = ""
    #     cells[1].text = ""

    # for row in table.rows:
    #     set_row_height(row, 20)

    return generate_admission_filename(data['FullName'], doc)

def send_whatsapp_message(media_url):
    message = client.messages.create(
        body="Here's the generated admission form:",
        from_=TWILIO_WHATSAPP_NUMBER,
        to=TARGET_WHATSAPP_NUMBER,
        media_url=[media_url]
    )
    print(f"Message SID: {message.sid}")

if __name__ == "__main__":
    app.run(debug=True)
